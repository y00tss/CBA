"""
Document processing module
"""
from datetime import datetime
import re
import uuid
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.shared import Pt
import json
import logging
from services.logger.logger import Logger
import os

logger = Logger(__name__, level=logging.INFO, log_to_file=True,
                filename='workflow.log').get_logger()

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


class DocumentWorkFlow:
    def __init__(self, path: str):
        self.path = path
        self.document = self._get_document()

        self.format_issues = []
        self.required_format_actions = []

        self.citation_issues = []
        self.required_citation_actions = []

    def _get_document(self):
        """Get document"""
        try:
            looking_for = os.path.join(BASE_DIR, self.path)
            return docx.Document(looking_for)
        except Exception as e:
            logger.error(f"Error getting document: {e}")

    async def start_flow(self):
        """Start document processing"""
        # 1. General Document Formatting Rules
        await self._front()
        await self._margins()
        await self._line_spacing()
        # await self._running_head()
        # await self._page_numbers()

        # 2. Document Structure Overview
        await self._title_page()
        await self._abstract()

        # 3. Keywords
        await self._keywords()

        # 4. Main Text
        await self._main_text()
        await self._in_text_citations()
        await self._heading_levels()
        # await self._tables()
        # await self._figures()

        return self.document

    async def create_report(self):
        """Create report on the issues found"""
        report = {
            "format_issues": {
                "issues": self.format_issues,
                "required_actions": self.required_format_actions
            },
            "citation_issues": {
                "issues": self.citation_issues,
                "required_actions": self.required_citation_actions
            }
        }
        return report

    async def _front(self):
        """Check front page"""
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if run.font.name != 'Times New Roman':
                    self.format_issues.append(f"Times New Roman was used for: '{run.text}'")

                    run.font.name = 'Times New Roman'

                if run.font.size and run.font.size.pt != 12:
                    self.format_issues.append(f"Font size 12 px was user for text: '{run.text}'")

                    run.font.size = docx.shared.Pt(12)

    async def _margins(self):
        """Check margins"""
        sections = self.document.sections
        for section in sections:
            if (section.left_margin.inches != 1 or
                    section.right_margin.inches != 1 or
                    section.top_margin.inches != 1 or
                    section.bottom_margin.inches != 1):
                section.left_margin = 1
                section.right_margin = 1
                section.top_margin = 1
                section.bottom_margin = 1
                self.format_issues.append("Margins were corrected to 1 inch on all sides")

    async def _line_spacing(self):
        """Check line spacing"""
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                if paragraph.paragraph_format.line_spacing != 2:
                    paragraph.paragraph_format.line_spacing = 2
                    self.format_issues.append(f"Line spacing corrected to 2 for paragraph: '{paragraph.text}'")

                space_after = paragraph.paragraph_format.space_after
                space_before = paragraph.paragraph_format.space_before

                if space_after is not None and space_after > 0:
                    paragraph.paragraph_format.space_after = 0
                    self.format_issues.append(f"Extra space after paragraph removed: '{paragraph.text}'")

                if space_before is not None and space_before > 0:
                    paragraph.paragraph_format.space_before = 0
                    self.format_issues.append(f"Extra space before paragraph removed: '{paragraph.text}'")

    async def _running_head(self):
        """Check running head"""
        header = self.document.sections[0].header

        if not header.paragraphs:
            paragraph = header.paragraphs.add_paragraph()
            paragraph.text = "RUNNING HEAD: TITLE OF THE PAPER"
            paragraph.paragraph_format.alignment = 0
            self.format_issues.append({"issue": "Running head added to header"})
        else:
            running_head = header.paragraphs[0].text
            if not running_head.isupper():
                header.paragraphs[0].text = running_head.upper()
                self.format_issues.append({"issue": "Running head corrected to uppercase", "location": running_head})

            if not header.paragraphs[0].alignment == 0:
                header.paragraphs[0].paragraph_format.alignment = 0
                self.format_issues.append({"issue": "Running head aligned to left", "location": running_head})

    async def _page_numbers(self):
        """Check and fix page numbers in header"""
        header = self.document.sections[0].header

        page_number = None
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if run.text.isdigit():
                    page_number = run.text

        if not page_number:
            paragraph = header.paragraphs.add()
            paragraph.alignment = 2
            paragraph.text = "1"
            self.format_issues.append({"issue": "Page number added to header, right-aligned"})
        else:
            if not header.paragraphs[-1].alignment == 2:
                header.paragraphs[-1].paragraph_format.alignment = 2
                self.format_issues.append({"issue": "Page number aligned to right"})

    async def _title_page(self):
        """Check title page"""
        first_page = self.document.paragraphs[:12]
        found_title = False
        for para in first_page:
            if para.style.name == 'Title':
                found_title = True
                if not self._is_title_case(para.text):
                    self.format_issues.append(f"Title case was used for title: {para.text}")
                    para.text = para.text.title()

                if not self._is_centered(para):
                    self.format_issues.append(f"Title was centered for title: {para.text}")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not any(run.bold for run in para.runs):
                    self.format_issues.append(f"Title was bolded for title: {para.text}")
                    for run in para.runs:
                        run.bold = True
                break

        if not found_title:
            self.required_format_actions.append("Add title to upper half of first page")

        author_info_found = False
        for para in first_page[1:]:
            if para.text.strip():
                author_info_found = True
                if not self._is_centered(para):
                    self.format_issues.append(f"Author information was centered")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if para.paragraph_format.line_spacing != Pt(24):
                    self.format_issues.append(f"Author information line spacing was corrected")
                    para.paragraph_format.line_spacing = Pt(24)
                break

        if not author_info_found:
            self.required_format_actions.append("Add author information to upper half of first page")

        author_note_found = False
        for para in first_page:
            if 'Author Note' in para.text:
                author_note_found = True
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.format_issues.append(f"Author Note was centered: {para.text}")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not any(run.bold for run in para.runs):
                    self.format_issues.append(f"Author Note was bolded for title: {para.text}")
                    for run in para.runs:
                        run.bold = True
                break

        if not author_note_found:
            self.required_format_actions.append("Add Author Note to upper half of first page")

    async def _abstract(self):
        """Check and correct Abstract section formatting"""
        found_abstract = False
        abstract_page = False
        for i, paragraph in enumerate(self.document.paragraphs):
            if paragraph.text.strip().lower() == "abstract":
                found_abstract = True

                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.format_issues.append(f"Abstract heading was centered: {paragraph.text}")
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if not any(run.bold for run in paragraph.runs):
                    self.format_issues.append("Abstract heading was bolded: {paragraph.text}")
                    for run in paragraph.runs:
                        run.bold = True

                if paragraph.page_break_before is None:
                    self.format_issues.append("Abstract was replaced on a separate page")
                    paragraph.insert_paragraph_before('\n', style='Normal').page_break_before = True
                abstract_page = True

                if len(self.document.paragraphs) > i + 1:
                    abstract_text = self.document.paragraphs[i + 1]

                    if abstract_text.paragraph_format.first_line_indent:
                        self.format_issues.append(f"Abstract first line indent removed: {abstract_text.text}")
                        abstract_text.paragraph_format.first_line_indent = None

                    word_count = len(abstract_text.text.split())
                    if word_count > 250:
                        self.format_issues.append("Abstract was cut to 250 words")
                        abstract_text.text = ' '.join(abstract_text.text.split()[:250])

                break

        if not found_abstract:
            self.required_format_actions.append("Abstract heading should be added to the document")
        elif not abstract_page:
            self.required_format_actions.append("Abstract heading should be on a separate page")

    async def _keywords(self):
        """Check and correct Keywords section formatting"""
        found_keywords = False
        for i, paragraph in enumerate(self.document.paragraphs):
            # Найти строку ниже Abstract
            if paragraph.text.strip().lower() == "abstract":
                if len(self.document.paragraphs) > i + 2:
                    keywords_paragraph = self.document.paragraphs[i + 2]
                    if not paragraph.text.lower().startswith("keywords:"):
                        self.required_format_actions.append("Keywords heading should begin with word: 'Keywords:'")
                    if not any(run.font.italic for run in paragraph.runs):
                        self.format_issues.append("Keywords heading is not italicized")
                        for run in paragraph.runs:
                            run.italic = True

                    keywords_paragraph.paragraph_format.left_indent = Pt(0.5 * 72)
                    self.format_issues.append("Keywords section indented by 0.5 inches")

                    keywords_text = keywords_paragraph.text.split(":")[1].strip()
                    keywords_lower = ', '.join([word.strip().lower() for word in keywords_text.split(",")])
                    keywords_paragraph.clear()
                    run = keywords_paragraph.add_run(f"Keywords: {keywords_lower}")
                    run.italic = True
                    found_keywords = True
                    break

        if not found_keywords:
            self.required_format_actions.append("Add Keywords section below Abstract")

    async def _main_text(self):
        """Check and format main text according to APA requirements"""
        self.document.add_section(WD_SECTION.NEW_PAGE)

        original_title = self.document.paragraphs[0].text
        title_paragraph = self.document.add_paragraph(original_title)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.runs[0].bold = True

        self.format_issues.append("Main text started on a new page with repeated title in bold.")

    async def _in_text_citations(self):
        """Check in-text citations for author-date format"""
        citation_pattern = r"\(([^()]+), (\d{4})(, p\. \d+)?\)"

        for paragraph in self.document.paragraphs:
            text = paragraph.text

            citations = re.findall(citation_pattern, text)

            if citations:
                for citation in citations:
                    author_part, year, page = citation
                    corrected_citation = f"({author_part.strip()}, {year}{page if page else ''})"

                    if corrected_citation not in text:
                        try:
                            paragraph.text = text.replace(f"({citation[0]}, {citation[1]}{citation[2]})",
                                                          corrected_citation)
                            self.format_issues.append("Corrected in-text citation format to author-date.")
                        except Exception as e:
                            logger.error(f"Error correcting in-text citation '{citation}': {e}")

    async def _heading_levels(self):
        """Format APA-style headings based on their levels"""
        for paragraph in self.document.paragraphs:
            if paragraph.style.name == 'Heading 1':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].bold = True
                paragraph.runs[0].text = await self._to_title_case(paragraph.text)

            elif paragraph.style.name == 'Heading 2':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.runs[0].bold = True
                paragraph.runs[0].text = await self._to_title_case(paragraph.text)

            elif paragraph.style.name == 'Heading 3':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.runs[0].bold = True
                paragraph.runs[0].italic = True
                paragraph.runs[0].text = await self._to_title_case(paragraph.text)

            elif paragraph.style.name == 'Heading 4':
                paragraph.paragraph_format.left_indent = Inches(0.5)
                paragraph.runs[0].bold = True
                paragraph.runs[0].text = await self._to_title_case(paragraph.text) + "."
                paragraph.runs[0].space_after = 0

            elif paragraph.style.name == 'Heading 5':
                paragraph.paragraph_format.left_indent = Inches(0.5)
                paragraph.runs[0].bold = True
                paragraph.runs[0].italic = True
                paragraph.runs[0].text = await self._to_title_case(paragraph.text) + "."
                paragraph.runs[0].space_after = 0

            self.format_issues.append(f"Formatted heading: '{paragraph.text}' to APA style.")

    async def _to_title_case(self, text):
        """Helper function to convert text to title case."""
        return ' '.join([word.capitalize() for word in text.split()])

    async def _tables(self):
        """Format tables according to APA style"""
        table_count = 1
        for table in self.document.tables:
            table_number = f"Table {table_count}"
            table_title_paragraph = self.document.add_paragraph()
            table_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_table_number = table_title_paragraph.add_run(table_number)
            run_table_number.bold = True

            run_title = table_title_paragraph.add_run(": Title")
            run_title.italic = True

            for row in table.rows:
                for cell in row.cells:
                    cell_borders = cell._element.findall(".//w:tcBorders", namespaces={"w": qn("w")})
                    if cell_borders:
                        for border in cell_borders:
                            top_border = border.find(qn("w:top"))
                            bottom_border = border.find(qn("w:bottom"))

                            if row == table.rows[0] and top_border is not None:
                                top_border.set("w:val", "single")
                            elif row == table.rows[-1] and bottom_border is not None:
                                bottom_border.set("w:val", "single")
                            else:
                                for side in ["left", "right"]:
                                    side_border = border.find(qn(f"w:{side}"))
                                    if side_border is not None:
                                        side_border.set("w:val", "nil")

            for cell in table.columns[0].cells:
                cell.text = cell.text.strip().title()

            # Обновляем текст вывода format_issues
            self.format_issues.append(f"Table {table_count} formatted according to APA style.")
            table_count += 1

    async def _figures(self):
        """Format figures according to APA style"""
        figure_count = 1
        for shape in self.document.inline_shapes:
            figure_paragraph = shape._inline.getparent().addnext(self.document.add_paragraph())

            figure_label = f"Figure {figure_count}"
            caption_paragraph = self.document.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run_label = caption_paragraph.add_run(figure_label)
            run_label.bold = True

            run_caption = caption_paragraph.add_run(": Brief description here")
            run_caption.italic = True
            run_caption.font.size = Pt(10)

            new_section = False
            if new_section:
                self.document.add_paragraph("\f")

            self.format_issues.append(f"{figure_label} formatted with APA-style caption.")
            figure_count += 1

    # """ENCAPSULATED FUNCTIONS"""
    def _is_title_case(self, text: str) -> bool:
        """Check if text is in title case"""
        words = text.split()
        for word in words:
            if not word[0].isupper():
                return False
        return True

    def _is_centered(self, paragraph) -> bool:
        """Check if paragraph is centered"""
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER

    async def get_updated_document(self, user_name: str):
        """Convert to docx after checking"""
        file_name = f"updated_document_{datetime.now().date()}_{uuid.uuid4()}.docx"
        file_path = f"articles/documents/{user_name}/{file_name}"

        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        self.document.save(file_path)

        return file_path
